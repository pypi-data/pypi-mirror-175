import docx

from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, PageBreak, Paragraph

STYLES = getSampleStyleSheet()


def create_test_docx(filename: str = 'document_to_convert.docx'):
    """Create dummy docx."""
    docx_document = docx.Document()
    docx_document.add_heading('The title of the document')
    docx_document.add_paragraph('Some interesting text.')
    docx_document.save(filename)


def create_test_pdf_single(filename: str = 'document_to_convert_single.pdf'):
    """Create dummy single page pdf."""
    pdf_document_single = SimpleDocTemplate(filename)
    story = [
        Paragraph(
            'The title of the document',
            STYLES['Heading1']
        ),
        Paragraph(
            'Some interesting text.',
            STYLES['Normal']
        )
    ]
    pdf_document_single.build(story)


def create_test_pdf_multi(filename: str = 'document_to_convert_multi.pdf'):
    pdf_document_multi = SimpleDocTemplate(filename)
    story = [
        Paragraph(
            'The title of the document',
            STYLES['Heading1']
        ),
        Paragraph(
            'A subtitle',
            STYLES['Heading2']
        ),
        Paragraph(
            'Some interesting text on the first page.',
            STYLES['Normal']
        ),
        PageBreak(),
        Paragraph(
            'Another subtitle',
            STYLES['Heading2']
        ),
        Paragraph(
            'Some interesting text on the second page.',
            STYLES['Normal']
        ),
        PageBreak(),
        Paragraph(
            'Another subtitle',
            STYLES['Heading2']
        ),
        Paragraph(
            'Some text on the third page.',
            STYLES['Normal']
        ),
    ]
    pdf_document_multi.build(story)
